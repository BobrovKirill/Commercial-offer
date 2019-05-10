import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches

file = 'task.docx'
doc = docx.Document(file)
print(doc.tables[1].rows[3].cells[1].text)
#
# table[№] номер таблицы rows[№] номер строки cells[№] номер столбца


#table = doc.add_table(1, 3)
#table.style = 'Table Grid'
#table.autofit
#hdr_cells = table.rows[0].cells
#hdr_cells[0].text = 'Qty'
#hdr_cells[1].text = 'Id'
#hdr_cells[2].text = 'Desc'
#doc.save(file)

val=[1,2,3]
print(val[2])

#from docx import Document
#from docx.shared import Inches
#
#document = Document('demo.docx')
#
#################################
#################################
#################################
#
#table = document.add_table(rows=1, cols=3)
#hdr_cells = table.rows[0].cells
#hdr_cells[0].text = 'Qty'
#hdr_cells[1].text = 'Id'
#hdr_cells[2].text = 'Desc'
#
#row_cells = table.add_row().cells
#row_cells[0].text = 'Str0'
#row_cells[1].text = 'Str1'
#row_cells[2].text = 'Str2'
#
#row2_cells = table.add_row().cells
#row2_cells[0].text = 'Str00'
#row2_cells[1].text = 'Str11'
#row2_cells[2].text = 'Str22'
#table.style = 'Table Grid'
#################################
#################################
#################################
#
#document.save('demo.docx')