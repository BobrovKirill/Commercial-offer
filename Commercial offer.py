from openpyxl import load_workbook
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT



        # считывание с excel


all_product = {}
wb = load_workbook(filename='test.xlsx', data_only=True)
ws = wb.active

for data in ws.values:
    if isinstance(data[0], int):
        all_product[data[1]] = data[2:]


        # поиск товара и добавление его в таблицу word


def request(key):
    if key in all_product:
        i = 1
        global sum_price
        sum_price = 0
        for key,val in all_product.items():
            data_cells = table.add_row().cells
            data_cells[0].text = str(i)
            paragraph = data_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            data_cells[1].text = key
            paragraph = data_cells[1].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            data_cells[2].text = str(val[0])
            paragraph = data_cells[2].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            data_cells[3].text = str(val[1])
            paragraph = data_cells[3].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            data_cells[4].text = str(val[2])
            paragraph = data_cells[4].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            data_cells[5].text = str(val[3])
            paragraph = data_cells[5].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sum_price += val[3]
            i += 1
    else:
        print('В списке нет данного товара')
    return sum_price

        # Открытие ворда и добавление стилей


file = 'task1.docx' # сделать так чтоб пользователь указывал название и путь шаблона
file_save = 'task_new.docx' # сделать так чтоб пользователь указывал название файла и путь
doc = docx.Document(file)
table = doc.tables[1]
table.style = 'Table Grid'
table.autofit = True
request('apple')

        #формирование двух последних строк с объединением столбцов


nums_cells = len(table.rows)
data_cells = table.add_row().cells
for i in range(2):
    merget = table.cell(nums_cells,i).merge(table.cell(nums_cells,i+1))
    merget = table.cell(nums_cells,i+3).merge(table.cell(nums_cells,i+4))
data_cells[0].text = 'Итого, рублей'
data_cells[3].text = str(sum_price) # !!!! Создать тут формулу суммы
paragraph = data_cells[0].paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph = data_cells[3].paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
data_cells = table.add_row().cells
for i in range(2):
    merget = table.cell(nums_cells+1,i).merge(table.cell(nums_cells+1,i+1))
    merget = table.cell(nums_cells+1,i+3).merge(table.cell(nums_cells+1,i+4))
data_cells[0].text = 'В том числе НДС, руб:'
data_cells[3].text = 'Без НДС*'
paragraph = data_cells[0].paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph = data_cells[3].paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(file_save)
