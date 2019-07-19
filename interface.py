import tkinter as tk
from tkinter import ttk
import os
from tkinter import filedialog
from tkinter import *
from openpyxl import load_workbook
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import  WD_STYLE_TYPE
from docx.shared import Pt


def readyng_excel():  # Считывание с excel файла
    fileexcel = Main.fileexcel
    all_product = {}
    wb = load_workbook(filename=fileexcel, data_only=True)
    ws = wb.active

    for data in ws.values:
        if isinstance(data[0], int):
            all_product[data[0]] = data[1:]
    Main.all_product = all_product


def dir_excel():  # Выбор пути до файла
    fileexcel = filedialog.askopenfilename(filetypes=[('Excel files', '*.xlsx')])
    path_excel.set(fileexcel)
    Main.fileexcel = fileexcel


def dir_word():  # Выбор пути до файла
    fileword = filedialog.askopenfilename(filetypes=[('Word files', '*.docx')])
    path_word.set(fileword)
    Main.fileword = fileword


class Main(tk.Frame):  # создание главного окна
    def __init__(self, root):
        super().__init__(root)
        self.init_main()

    def init_main(self):
        toolbar = tk.Frame(bg='#d7d8e0', bd=2)
        toolbar.place(width=430, height=150)
        btn_open_excel = tk.Button(toolbar, text='Укажите путь до Excel файла', command=self.open_exel, bg='#d7d8e0',
                                   bd=2, compound=tk.TOP)
        btn_open_excel.place(x=30, y=10)
        btn_open_word = tk.Button(toolbar, text='Укажите путь до Word файла', command=self.open_word, bg='#d7d8e0',
                                  bd=2, compound=tk.TOP)
        btn_open_word.place(x=220, y=10)
        btn_save = tk.Button(toolbar, text='Укажите путь нового файла', command=path_save, bg='#d7d8e0',bd=2, compound=tk.TOP)
        btn_save.place(x=175, y=50)
        btn_start = tk.Button(toolbar, text='Создать файл', command=start, bg='#d7d8e0',bd=2, compound=tk.TOP)
        btn_start.place(x=175, y=100)

    def open_exel(self):
        child_excel()
    def open_word(self):
        child_word()


class child_excel(tk.Toplevel):  # создание окна для выбора файла
    def __init__(self):
        super().__init__(root)
        self.init_child()

    def init_child(self):
        self.title('Выбор Exel файла')
        self.geometry('550x150+700+350')
        self.resizable(False, False)
        self.grab_set()
        self.focus_set()

        labal = ttk.Label(self, text='Выберете путь к Excel файлу:')
        labal.place(x=10, y=30)

        self.entry_excel = ttk.Entry(self, width=45, textvariable=path_excel)
        self.entry_excel.place(x=180, y=30)

        btn_cancel = ttk.Button(self, text='Закрыть', command=self.destroy)
        btn_cancel.place(x=465, y=100)

        btn_ok = ttk.Button(self, text='Ok', command=self.button_excel)
        btn_ok.place(x=385, y=100)
        btn_ok.bind('<Button-1>')

        btn_browse = ttk.Button(self, text='Обзор', command=dir_excel)
        btn_browse.place(x=465, y=28)
        btn_browse.bind('<Button-1>')

    def button_excel(self):
        readyng_excel()
        self.destroy()


class child_word(tk.Toplevel):  # создание окна для выбора файла
    def __init__(self):
        super().__init__(root)
        self.init_child()

    def init_child(self):
        self.title('Выбор Word файла')
        self.geometry('550x150+700+350')
        self.resizable(False, False)
        self.grab_set()
        self.focus_set()

        labal = ttk.Label(self, text='Выберете путь к Word файлу:')
        labal.place(x=10, y=30)

        self.entry_excel = ttk.Entry(self, width=45, textvariable=path_word)
        self.entry_excel.place(x=180, y=30)

        btn_cancel = ttk.Button(self, text='Закрыть', command=self.destroy)
        btn_cancel.place(x=465, y=100)

        btn_ok = ttk.Button(self, text='Ok', command=self.button_word)
        btn_ok.place(x=385, y=100)
        btn_ok.bind('<Button-1>')

        btn_browse = ttk.Button(self, text='Обзор', command=dir_word)
        btn_browse.place(x=465, y=28)
        btn_browse.bind('<Button-1>')

    def button_word(self):
        readyng_word()
        self.destroy()

        # поиск товара и добавление его в таблицу word


def readyng_word():
    def request():
        all_product = Main.all_product
        sum_price, val = 0, 0
        style = doc.styles.add_style('my_style', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        for key, value in all_product.items():
            data_cells = table.add_row().cells
            data_cells[0].text = str(key)
            i = 0
            for val in value:
                data_cells[i + 1].text = str(val)
                i += 1
            for j in range(len(value)+1):
                paragraph = data_cells[j].paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.style = 'my_style'
            sum_price += val

        # формирование двух последних строк с объединением столбцов

        nums_cells = len(table.rows)
        data_cells = table.add_row().cells
        for i in range(2):
            merget = table.cell(nums_cells, i).merge(table.cell(nums_cells, i + 1))
            merget = table.cell(nums_cells, i + 3).merge(table.cell(nums_cells, i + 4))
        data_cells[0].text = 'Итого, рублей'
        data_cells[3].text = str(sum_price)
        paragraph = data_cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph = data_cells[3].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        data_cells = table.add_row().cells
        for i in range(2):
            merget = table.cell(nums_cells + 1, i).merge(table.cell(nums_cells + 1, i + 1))
            merget = table.cell(nums_cells + 1, i + 3).merge(table.cell(nums_cells + 1, i + 4))
        data_cells[0].text = 'В том числе НДС, руб:'
        data_cells[3].text = 'Без НДС*'
        paragraph = data_cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph = data_cells[3].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Открытие ворда
    fileword = Main.fileword
    file = fileword
    doc = docx.Document(file)
    table = doc.tables[1]
    request()
    Main.doc = doc

def start():
    doc = Main.doc
    new_path = Main.new_path
    doc.save(new_path + '/Новое коммерч.docx')


def path_save():
    new_path = filedialog.askdirectory()
    Main.new_path = new_path

if __name__ == '__main__':
    root = tk.Tk()
    path_excel = StringVar()
    path_word = StringVar()
    app = Main(root)
    app.pack()
    root.title('Commercial Offer')
    root.geometry('430x150+700+350')
    root.resizable(False, False)
    root.mainloop()

