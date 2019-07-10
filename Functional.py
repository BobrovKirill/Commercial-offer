from openpyxl import load_workbook
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import  WD_STYLE_TYPE
from docx.shared import Pt


        # считывание с excel


all_product = {}
wb = load_workbook(filename='товары.xlsx', data_only=True)
ws = wb.active

for data in ws.values:
    if isinstance(data[0], int):
        all_product[data[0]] = data[1:]


        # поиск товара и добавление его в таблицу word


def request():
    global sum_price
    sum_price = 0
    val = 0
    style = doc.styles.add_style('my_style', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    for key,value in all_product.items():
        data_cells = table.add_row().cells
        data_cells[0].text = str(key)
        i = 0
        for val in value:
            data_cells[i+1].text = str(val)
            i+=1
        for j in range(len(value)+1):
            paragraph = data_cells[j].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.style = 'my_style'
        sum_price += val

        # Открытие ворда и добавление стилей


file = 'шаблон.docx' # сделать так чтоб пользователь указывал название и путь шаблона
file_save = 'Новое коммерческое.docx' # сделать так чтоб пользователь указывал название файла и путь
doc = docx.Document(file)
table = doc.tables[1]
request()

        #формирование двух последних строк с объединением столбцов


nums_cells = len(table.rows)
data_cells = table.add_row().cells
for i in range(2):
    merget = table.cell(nums_cells,i).merge(table.cell(nums_cells,i+1))
    merget = table.cell(nums_cells,i+3).merge(table.cell(nums_cells,i+4))
data_cells[0].text = 'Итого, рублей'
data_cells[3].text = str(sum_price)
paragraph = data_cells[0].paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph = data_cells[3].paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
data_cells = table.add_row().cells
for i in range(2):
    merget = table.cell(nums_cells+1,i).merge(table.cell(nums_cells+1,i+1))
    merget = table.cell(nums_cells+1,i+3).merge(table.cell(nums_cells+1,i+4))
data_cells[0].text = 'В том числе НДС, руб:'
data_cells[3].text = 'Без НДС*'
paragraph = data_cells[0].paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph = data_cells[3].paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(file_save)